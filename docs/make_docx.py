"""Erzeugt docs/PiLiDAR_Technisches_Konzept.docx (Architektur, Mess-Mathematik,
Montage, Verkabelung, Kalibrierung, QA, Inbetriebnahme) inkl. Bilder.
Aufruf: python docs/make_docx.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(__file__)
IMG = os.path.join(HERE, "images")
MOUNT_IMG = os.path.join(os.path.dirname(HERE), "hardware", "stl27l_mount.png")
OUT = os.path.join(HERE, "PiLiDAR_Technisches_Konzept.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def img(doc, path, width=6.3, caption=""):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].italic = True
            c.runs[0].font.size = Pt(9)
    else:
        para(doc, f"[Bild fehlt: {path}]", italic=True)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def build():
    doc = Document()
    # Titel
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PiLiDAR 2.0")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = ACCENT
    sub = doc.add_paragraph("Technisches Konzept — Web-Steuerung, USB-LiDAR (STL27L), "
                            "kontinuierliche Drehung")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    meta = doc.add_paragraph("Optimierte Neuumsetzung des Open-Source-Projekts PiLiDAR · "
                             "Raspberry Pi 4 · NEMA17/A4988(TMC2209) · STL27L")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].italic = True
    doc.add_page_break()

    # 1 Überblick
    h(doc, "1  Überblick & Ziele", 1)
    para(doc, "Diese Umsetzung baut den 360°-3D-Scanner PiLiDAR um auf eine schlanke "
              "Raspberry-Pi-Erfassung mit Web-Steuerung. Der Pi erfasst nur Rohdaten und "
              "steuert die Mechanik; die rechenintensive Punktwolken-Berechnung und die "
              "Visualisierung laufen im Browser des Clients — ohne jede Zusatzinstallation.")
    bullets(doc, [
        "Keine Hardware-Knöpfe — Steuerung vollständig über Webserver.",
        "LiDAR über USB (STL27L-Controllerboard, CP2102 → /dev/ttyUSB0).",
        "Rechenlast am Client (Browser, WebGL) — Pi liefert nur Roh-Frames.",
        "Zweite Messtechnik: kontinuierliche, konstante Drehung (ruckelfrei) neben der bestehenden Schrittmessung.",
        "Pro Scan ein eigener Ordner mit Rohdaten + Punktwolken in mehreren Formaten.",
        "Automatische Referenzfahrt nach jedem Scan, Kalibrierprogramm, Qualitätssicherung.",
        "IMU und Kamera sind vorbereitet, aber bewusst noch nicht umgesetzt.",
    ])

    h(doc, "2  Lizenz & Attribution", 1)
    para(doc, "Basis ist das Projekt PiLiDAR (github.com/PiLiDAR) unter CC BY-NC-SA 4.0. "
              "Diese abgeleitete Arbeit übernimmt dieselbe Lizenz: nicht-kommerzielle Nutzung "
              "mit Namensnennung und Weitergabe unter gleichen Bedingungen. Ein kommerzieller "
              "Geräteverkauf ist ohne ausdrückliche Genehmigung des Originalautors nicht zulässig.")

    # 3 Architektur
    h(doc, "3  Architektur", 1)
    para(doc, "Dünner Pi-Server (FastAPI) + dicker Web-Client. Der Pi liest die serielle "
              "LiDAR-Schnittstelle in einem eigenen Thread, validiert jedes Paket per CRC8 und "
              "streamt die Roh-Frames per WebSocket. Der Client rechnet Polar→Kartesisch und die "
              "3D-Revolve in einem Web Worker und rendert per WebGL.")
    h(doc, "3.1  Web-API", 2)
    table(doc, ["Endpunkt", "Methode", "Funktion"], [
        ["/api/status", "GET", "Zustand, Winkel, Paketrate, CRC-Rate"],
        ["/api/lidar/start | /stop", "POST", "Lidar-Only-Modus (2D-Live, Motor aus)"],
        ["/api/scan/start | /stop", "POST", "Scan starten (Modus A/B) / abbrechen"],
        ["/api/scans, /scans/{id}/download", "GET", "Scan-Liste / ZIP-Download"],
        ["/api/calibrate/rotation | /offset", "POST", "Kalibrierläufe"],
        ["/api/config/apply", "POST", "Kalibrierergebnisse in config.json übernehmen"],
        ["/ws/scan", "WS", "binärnaher Roh-Frame-Stream (init + frames)"],
    ])
    para(doc, "Keine Client-Installation: Alle JS-Bibliotheken werden vom Pi ausgeliefert; der "
              "3D-Viewer ist ein eigener, schlanker WebGL-Renderer (kein three.js, kein CDN).")

    # 4 Hardware & Verkabelung
    h(doc, "4  Hardware & Verkabelung", 1)
    para(doc, "Zwei Treibervarianten sind dokumentiert: A4988 (Ist-Zustand) und TMC2209 "
              "(leiseres/glatteres Upgrade für die kontinuierliche Drehung). Kernregeln:")
    bullets(doc, [
        "100 µF Elektrolytkondensator ZWINGEND direkt an VMOT/GND jedes Treibers — sonst zerstören LC-Spannungsspitzen den Treiber.",
        "0,1 µF Keramik (X7R) an VDD/Logik; getrennte Netzteile für Pi (5 V) und Motor (12 V); EINE gemeinsame Sternmasse.",
        "Motorphasen 18–20 AWG verdrillt, Signalleitungen 24–28 AWG; Schirm einseitig auf GND; Ferrit auf USB/Signal.",
        "LiDAR über USB — keine GPIO-UART-Pegelwandlung nötig (3,3-V-Logik bleibt geschützt).",
    ])
    para(doc, "Übersichtspläne:")
    img(doc, os.path.join(IMG, "verkabelung_a4988.png"), caption="Übersicht A4988 (Ist-Zustand)")
    img(doc, os.path.join(IMG, "verkabelung_tmc2209.png"), caption="Übersicht TMC2209 (Upgrade)")
    para(doc, "Detailgetreue, pin-genaue Pläne (komplette 40-Pin-Leiste, voller Treiber-Pinout, "
              "Kondensatoren an den korrekten Knoten, Spulenpaare, Sternmasse):")
    img(doc, os.path.join(IMG, "verkabelung_a4988_detail.png"), caption="Detailplan A4988 — pin-genau")
    img(doc, os.path.join(IMG, "verkabelung_tmc2209_detail.png"), caption="Detailplan TMC2209 — pin-genau (inkl. UART-Option)")
    img(doc, os.path.join(IMG, "gpio_pinout.png"), width=5.5, caption="Belegte GPIO-Pins (BCM)")
    para(doc, "Detaillierte Stückliste, Specs, GPIO-Belegung, Kondensatoren/Kabel und der "
              "Treibervergleich befinden sich in der Datei PiLiDAR_BOM_Specs.xlsx.")

    # 5 Montage
    h(doc, "5  Korrekte LiDAR-Montage", 1)
    para(doc, "Die gesamte Wolken-Geometrie hängt von der Einbaulage ab. Der STL27L wird AUF DER "
              "SEITE montiert: seine Bodenplatte steht senkrecht an der Rückplatte, die Spinachse "
              "zeigt horizontal nach außen (Y). Dadurch ist die Scan-Ebene VERTIKAL (X-Z) und kann "
              "vom Stepper um die senkrechte Z-Achse revolviert werden.")
    bullets(doc, [
        "OBEN/UNTEN: Bodenplatte senkrecht an der Rückplatte; Optikfenster (Band um den Kopf) rundum frei.",
        "Spinachse horizontal (Y); Scan-Ebene vertikal (X-Z).",
        "Stecker ZH1.5T-4P (Tx, PWM, GND, VCC) unten; Kabel mitdrehend mit Zugentlastung.",
        "Läge der LiDAR flach, wäre die Scan-Ebene horizontal — es entstünde KEINE 3D-Wolke.",
    ])
    img(doc, os.path.join(IMG, "lidar_montage.png"), caption="Einbaulage und Ausrichtung des STL27L")

    # 6 Messtechnik & Mathematik
    h(doc, "6  Messtechnik & Punkt-Berechnung", 1)
    para(doc, "Pro Messpunkt liefert der LiDAR einen Winkel α (in der Scan-Ebene) und eine Distanz r. "
              "Die Umrechnung in 3D ist identisch zum bewährten PiLiDAR und für beide Messmodi gleich:")
    bullets(doc, [
        "1) Ebene:  x = r·cos α,  z = r·sin α,  y = 0  (vertikale X-Z-Ebene).",
        "2) Y-Rotation um angle_offset (mechanische Kippkorrektur, hier −1,05°).",
        "3) + Positionsversatz (0, MODEL_Y=−37,5 mm, MODEL_Z=−41,9 mm) — Sensor außerhalb der Drehachse.",
        "4) Z-Rotation um −z_angle (Revolve um die senkrechte Drehachse).",
    ])
    img(doc, os.path.join(IMG, "scan_vorgang.png"), caption="Scan-Vorgang: Aufbau, Offset, Punkt-Berechnung")

    h(doc, "6.1  Koordinatensystem & Offsets (technische Zeichnung)", 2)
    para(doc, "Es gilt ein rechtshändiges Koordinatensystem (X×Y=Z) mit Ursprung O auf der "
              "Drehachse: +Z senkrecht nach oben (= Drehachse, Revolve-Achse), +Y horizontal "
              "entlang der LiDAR-Spinachse, +X quer dazu. Maße in mm, Winkel in Grad.")
    img(doc, os.path.join(IMG, "geometrie_koordinaten.png"),
        caption="Koordinatensystem & bemaßte Offsets (Ansicht A: Y-Z-Ebene, Ansicht B: Scan-Ebene X-Z)")

    h(doc, "6.1.1  angle_offset — präzise Bedeutung", 3)
    para(doc, "angle_offset ist eine Drehung der Messpunkte um die Y-Achse — und Y ist die Normale "
              "der vertikalen Scan-Ebene (X-Z). Da die Punkte in dieser Ebene liegen, ist eine "
              "Drehung um Y mathematisch IDENTISCH mit einer Verschiebung des Nullwinkels jedes "
              "Messpunkts INNERHALB der Ebene:")
    para(doc, "Rot_Y(φ)·(r·cos α, 0, r·sin α) = (r·cos(α−φ), 0, r·sin(α−φ))   ⇒   α → α − φ", italic=True)
    para(doc, "Physikalisch korrigiert angle_offset damit die rotatorische Einbau-„Clockung“ des "
              "LiDAR um seine eigene Spinachse (Y): den Versatz zwischen dem internen 0°-Bezug des "
              "Sensors und der gewünschten Null der Scan-Ebene. Es ist KEINE Verkippung aus der "
              "Ebene heraus (das wäre eine Drehung um X oder Z), sondern eine reine "
              "In-Ebenen-Winkelkorrektur. Typwert hier: −1,05°.")

    h(doc, "6.1.2  Warum die Offsets in Achsrichtung negativ sind", 3)
    para(doc, "position_offset = (0, MODEL_Y, MODEL_Z) ist die Lage des LiDAR-Optikzentrums S "
              "relativ zum Ursprung O. Das Vorzeichen folgt aus der Wahl des Koordinatensystems, "
              "nicht aus einer physikalischen Notwendigkeit:")
    bullets(doc, [
        "MODEL_Y = −37,5 mm: S liegt auf der −Y-Seite der Drehachse (gegen die gewählte +Y-Richtung der Spinachse).",
        "MODEL_Z = −41,9 mm: S liegt baulich UNTERHALB des Datums O (gegen +Z nach oben).",
        "Die Translation wird VOR der Z-Revolve angewandt, damit die Drehachse durch den Ursprung verläuft — sie verschiebt die Sensorpunkte auf die Achse.",
        "Bei umgekehrter +Y/+Z-Definition oder anderem Datum kehrten sich die Vorzeichen entsprechend um.",
    ])
    para(doc, "Wichtige Eigenschaft von MODEL_Z: Da die Revolve um Z die Z-Komponente nicht "
              "verändert, wirkt ein MODEL_Z-Fehler als reine vertikale Verschiebung der GESAMTEN "
              "Wolke — er verzerrt die Geometrie NICHT, sondern legt nur das Höhen-Datum fest. "
              "MODEL_Y dagegen ist der Hebelarm der Revolve und beeinflusst die Form unmittelbar.")

    h(doc, "6.2  Modus A — schrittweise (Original)", 2)
    para(doc, "Der Motor bewegt sich in diskreten Schritten mit kurzer Setzpause; z_angle ist je "
              "Schritt konstant. Folge: ruckelige Bewegung, Setzpausen, geringe Vibrationen.")
    h(doc, "6.3  Modus B — kontinuierliche konstante Drehung (neu)", 2)
    para(doc, "Der Motor dreht mit konstanter Winkelgeschwindigkeit ω. Erzeugt wird sie jitterfrei "
              "über Hardware-PWM am STEP-Pin (GPIO19 = PWM-Kanal 1):")
    para(doc, "ω = f_PWM / (microsteps · gear_ratio) · step_angle   [°/s]", italic=True)
    para(doc, "Jedem Paket wird ein kontinuierlicher Winkel zugeordnet:")
    para(doc, "z_angle = ω · (t_Paket − t_Start)   (Monotonic-Clock; Gegenprobe über LiDAR-Timestamp)", italic=True)
    para(doc, "Da ein Paket nur ~0,55 ms dauert (bei ω≈6°/s also ~0,003° Drift), genügt ein Winkel "
              "pro Paket. Die Revolve-Mathematik bleibt unverändert — es ändert sich nur die "
              "Feinheit und die Quelle von z_angle. Ergebnis: lückenlose, ruckelfreie Abtastung "
              "ohne Setzpausen.")
    para(doc, "Drehwinkel etwas > 180°: Wegen des Sensor-Offsets treffen die Strahlen bei φ und "
              "φ+180° parallel versetzt NICHT exakt dieselben Flächen. Erst ein Zuschlag Δ "
              "(overlap_deg) sorgt für lückenlose Vollabdeckung UND echte Doppelerfassung (Basis "
              "für QA/Kalibrierung).")

    # 7 Kalibrierung
    h(doc, "7  Kalibrierprogramm", 1)
    para(doc, "Web-geführt, nur mit LiDAR-Eigendaten (kein Spezialwerkzeug). Ergebnisse können "
              "per Klick in config.json übernommen werden.")
    bullets(doc, [
        "Rotation/Getriebeverhältnis: 360°-Lauf, Kreuzkorrelation der doppelt gesehenen Profile → echter Winkel pro Schritt.",
        "Offset (MODEL_Y/Z) + angle_offset: Selbstkonsistenz-Optimierung in der Überlappungszone (>180°); schätzt Δ mit.",
        "Ausrichtung: geführter manueller Wand-Check zur Verifikation der vertikalen Scan-Ebene.",
    ])
    h(doc, "7.1  Automatische Ermittelbarkeit & Zuverlässigkeit", 2)
    para(doc, "Können ALLE Offsets automatisch durch eine Kalibrierfahrt bestimmt werden? — "
              "Größtenteils ja, mit einer wichtigen Ausnahme (MODEL_Z). Entscheidend ist die "
              "Beobachtbarkeit aus den Eigendaten:")
    table(doc, ["Parameter", "Auto-Kalibrierung", "Beobachtbarkeit / Methode", "Zuverlässigkeit"], [
        ["gear_ratio (Schritte/Grad)", "ja", "360°-Lauf, Kreuzkorrelation der Profile", "hoch in strukturierten Räumen; schwach bei symmetrischen/leeren Szenen"],
        ["angle_offset (Clocking)", "ja", "Selbstkonsistenz φ/φ+180° (rotatorischer Versatz)", "hoch; Nahbereichsstruktur trennt es von MODEL_Y"],
        ["MODEL_Y (Hebelarm)", "ja", "Parallaxe in der Überlappung skaliert mit MODEL_Y", "mittel–hoch; braucht variierte Distanzen"],
        ["MODEL_Z (Höhen-Datum)", "nein", "nicht beobachtbar — nur globaler vertikaler Versatz, keine Verzerrung", "nicht nötig; aus Mechanik/externer Referenz (z.B. Boden) setzen"],
        ["overlap Δ (>180°)", "ja", "wird in der Offset-Optimierung mitgeschätzt", "gut"],
    ])
    para(doc, "Voraussetzungen für verlässliche Ergebnisse: Szene mit Struktur (keine leeren/"
              "symmetrischen Räume), variierte Distanzen (Nah- und Fernbereich), ausreichend "
              "Überlappungspunkte, KEIN Schrittverlust während des Kalibrierscans (sonst wird die "
              "Annahme konstanter ω verletzt). Der Optimierer wird mit den mechanischen "
              "Nennwerten gestartet, um lokale Minima zu vermeiden. angle_offset und MODEL_Y "
              "können im Fernbereich leicht korrelieren (eine kleine Clockung sieht aus wie ein "
              "kleiner seitlicher Versatz) — Nahbereichsstruktur löst diese Mehrdeutigkeit. "
              "Das Loop-Closure-Residuum (AP/QA) dient zugleich als Gütemaß: sinkt es nach der "
              "Kalibrierung deutlich, sind die Parameter konsistent.")

    # 8 QA
    h(doc, "8  Fehlererkennung & Qualitätssicherung", 1)
    para(doc, "Pro Scan wird automatisch ein QA-Bericht erstellt (in meta.json, Ampelstatus in der UI):")
    table(doc, ["Metrik", "Bedeutung"], [
        ["CRC8-Fehlerrate", "Verkabelung/Baudrate/EMV-Probleme"],
        ["Paketrate & Drops", "Serial-Overrun/USB-Probleme/Pufferüberlauf"],
        ["LiDAR-RPM-Stabilität", "blockierter/instabiler LiDAR-Motor"],
        ["Loop-Closure-Residuum", "Schrittverlust, Schlupf, Vibration, Fehlkalibrierung"],
        ["Winkel-Drift (Modus B)", "PWM-Winkel vs. LiDAR-Timestamp"],
        ["Homing-Kontrolle", "Rotor zurück in Ausgangsposition (≈0°)"],
    ])

    # 9 Datenablage
    h(doc, "9  Datenablage (ein Ordner pro Scan)", 1)
    para(doc, "scans/<id>/ enthält meta.json, raw/ (Rohpakete + Winkel/Zeit, .bin/.npz/.jsonl), "
              "pointcloud/ (.ply/.xyz, optional .e57) sowie vorbereitete images/ und sensors/ für "
              "spätere Kamera-/IMU-Daten. Download als ZIP über die Web-UI.")

    # 10 Inbetriebnahme
    h(doc, "10  Inbetriebnahme", 1)
    bullets(doc, [
        "Abhängigkeiten: pip install -r requirements.txt (auf dem Pi zusätzlich rpi-lgpio, rpi-hardware-pwm).",
        "Hardware-PWM aktivieren: in /boot/firmware/config.txt 'dtoverlay=pwm-2chan' eintragen, neu starten.",
        "LiDAR-USB-Rechte: Nutzer in Gruppe 'dialout' (sudo usermod -a -G dialout <user>).",
        "Start: python -m backend.app  (Server auf Port 8000).",
        "Demo ohne Hardware: PILIDAR_MOCK=1 python -m backend.app  (synthetischer LiDAR).",
        "Bedienung: Browser auf http://<pi-ip>:8000 — 2D-Live, Scan (Modus A/B), Kalibrierung, Export.",
    ])

    # 11 Halter
    h(doc, "11  STL27L-Halter (3D-Druck)", 1)
    para(doc, "Parametrischer Halter (hardware/stl27l_mount.scad) für die korrekte Seitenmontage. "
              "Maße aus dem STL-27L-Datenblatt (54,00 × 46,29 × 34,80 mm). Klemmung nur am Sockel, "
              "Optikfenster frei; geringe Masse für die kontinuierliche Drehung. Druck in PETG.")
    img(doc, MOUNT_IMG, width=4.8, caption="Parametrischer STL27L-Halter (Seitenmontage)")

    doc.save(OUT)
    print("ok", OUT)


if __name__ == "__main__":
    build()
